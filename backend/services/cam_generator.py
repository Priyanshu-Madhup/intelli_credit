"""
CAM Generator — builds a professional Credit Appraisal Memorandum (.docx)
"""
import io
from datetime import datetime
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_NAVY  = RGBColor(0x1E, 0x29, 0x3B)
BLUE       = RGBColor(0x25, 0x63, 0xEB)
MID_GRAY   = RGBColor(0x64, 0x74, 0x8B)
RED        = RGBColor(0xDC, 0x26, 0x26)
AMBER      = RGBColor(0xD9, 0x77, 0x06)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.color.rgb = BLUE if level == 1 else DARK_NAVY
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '2')
        bot.set(qn('w:color'), '2563EB')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def _para(doc, text, size=10, color=None, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def _kv_table(doc, rows):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (label, value) in enumerate(rows):
        row = tbl.rows[i]
        key_cell = row.cells[0]
        _set_cell_bg(key_cell, 'F1F5F9')
        kp = key_cell.paragraphs[0]
        krun = kp.add_run(label)
        krun.bold = True
        krun.font.size = Pt(9)
        krun.font.color.rgb = MID_GRAY
        val_cell = row.cells[1]
        vp = val_cell.paragraphs[0]
        vrun = vp.add_run(str(value) if value is not None else '—')
        vrun.font.size = Pt(10)
    doc.add_paragraph()

def generate_cam_docx(assessment: Dict[str, Any]) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    company_name = assessment.get('company_name', 'Unknown')
    sector       = assessment.get('sector', '—')
    decision     = assessment.get('decision', '—').upper()
    risk_score   = assessment.get('risk_score', '—')
    rec_loan     = assessment.get('recommended_loan_cr')
    req_loan     = assessment.get('requested_loan_cr')
    interest     = assessment.get('interest_rate_pct')
    tenor        = assessment.get('tenor_months')
    conditions   = assessment.get('conditions', [])
    alerts       = assessment.get('risk_alerts', [])
    reasoning    = assessment.get('reasoning', '')
    fo           = assessment.get('financial_overview', {})
    sb           = assessment.get('score_breakdown', {})
    yearly_trend = assessment.get('yearly_trend', [])
    report_date  = datetime.now().strftime('%d %b %Y')
    report_time  = datetime.now().strftime('%H:%M IST')

    # TITLE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    t1 = title_p.add_run('CREDIT APPRAISAL MEMORANDUM')
    t1.bold = True
    t1.font.size = Pt(18)
    t1.font.color.rgb = DARK_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(12)
    s1 = sub_p.add_run(f'{company_name}  |  {sector}  |  {report_date}')
    s1.font.size = Pt(10)
    s1.font.color.rgb = MID_GRAY
    s1.italic = True

    # HR line
    hr_p = doc.add_paragraph()
    hr_p.paragraph_format.space_after = Pt(12)
    pPr = hr_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2563EB')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Decision banner
    decision_colors = {
        'APPROVED':    ('D1FAE5', '065F46'),
        'CONDITIONAL': ('FEF3C7', '92400E'),
        'REJECTED':    ('FEE2E2', '991B1B'),
    }
    bg_hex, fg_hex = decision_colors.get(decision, ('F1F5F9', '1E293B'))
    banner_tbl = doc.add_table(rows=1, cols=3)
    banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate([
        ('DECISION',    decision),
        ('RISK SCORE',  f'{risk_score} / 100'),
        ('RECOMMENDED', f'Rs.{rec_loan} Cr' if rec_loan is not None else '—'),
    ]):
        cell = banner_tbl.rows[0].cells[i]
        _set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        lrun = p.add_run(lbl + '\n')
        lrun.font.size = Pt(8)
        lrun.bold = True
        lrun.font.color.rgb = RGBColor.from_string(fg_hex)
        vrun = p.add_run(val)
        vrun.font.size = Pt(13)
        vrun.bold = True
        vrun.font.color.rgb = RGBColor.from_string(fg_hex)
    doc.add_paragraph()

    # SECTION 1 — CHARACTER
    _heading(doc, '1. CHARACTER — Company Background')
    _kv_table(doc, [
        ('Company Name',     company_name),
        ('Sector',           sector),
        ('Loan Requested',   f'Rs.{req_loan} Cr' if req_loan is not None else '—'),
        ('Loan Recommended', f'Rs.{rec_loan} Cr' if rec_loan is not None else '—'),
        ('Decision',         decision.capitalize()),
        ('Report Date',      report_date),
    ])

    # SECTION 2 — CAPACITY
    _heading(doc, '2. CAPACITY — Financial Analysis')
    _kv_table(doc, [
        ('Annual Revenue',      fo.get('annual_revenue', '—')),
        ('Revenue Trend',       f"{fo.get('annual_revenue_delta','—')} ({fo.get('annual_revenue_trend','—')})"),
        ('Net Profit',          fo.get('net_profit', '—')),
        ('Net Profit Margin',   fo.get('net_profit_margin', '—')),
        ('Total Debt',          fo.get('total_debt', '—')),
        ('Debt / Equity Ratio', fo.get('de_ratio', '—')),
        ('GST Turnover',        fo.get('gst_turnover', '—')),
    ])

    if yearly_trend:
        _heading(doc, 'Year-on-Year Financial Trend', level=2)
        trend_tbl = doc.add_table(rows=1 + len(yearly_trend), cols=4)
        trend_tbl.style = 'Table Grid'
        for j, h in enumerate(['Year', 'Revenue (Cr)', 'Profit (Cr)', 'Debt (Cr)']):
            cell = trend_tbl.rows[0].cells[j]
            _set_cell_bg(cell, '1E3A5F')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, rd in enumerate(yearly_trend):
            row = trend_tbl.rows[i + 1]
            bg = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
            for j, val in enumerate([str(rd.get('year','—')), str(rd.get('revenue','—')), str(rd.get('profit','—')), str(rd.get('debt','—'))]):
                cell = row.cells[j]
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    if sb:
        _heading(doc, 'Credit Score Breakdown', level=2)
        score_rows = [
            ('Financial Health',    sb.get('financial_health',    {}).get('score', 0), 30),
            ('Repayment History',   sb.get('repayment_history',   {}).get('score', 0), 25),
            ('Collateral Coverage', sb.get('collateral_coverage', {}).get('score', 0), 20),
            ('Management Quality',  sb.get('management_quality',  {}).get('score', 0), 15),
            ('Market Position',     sb.get('market_position',     {}).get('score', 0), 10),
        ]
        score_tbl = doc.add_table(rows=1 + len(score_rows), cols=3)
        score_tbl.style = 'Table Grid'
        for j, h in enumerate(['Dimension', 'Weight', 'Score / 100']):
            cell = score_tbl.rows[0].cells[j]
            _set_cell_bg(cell, '1E3A5F')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
        for i, (dim, score, weight) in enumerate(score_rows):
            row = score_tbl.rows[i + 1]
            bg = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
            for j, val in enumerate([dim, f'{weight}%', f'{score}']):
                cell = row.cells[j]
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(10)
                if j == 2:
                    run.bold = True
                    run.font.color.rgb = GREEN if score >= 70 else (AMBER if score >= 40 else RED)
        doc.add_paragraph()

    # SECTION 3 — CAPITAL & COLLATERAL
    _heading(doc, '3. CAPITAL & COLLATERAL — Risk Assessment')
    if alerts:
        for alert in alerts:
            sev = alert.get('severity', 'low').lower()
            col = RED if sev == 'high' else (AMBER if sev == 'medium' else GREEN)
            badge = {'high': '[WARN HIGH]', 'medium': '[WARN MEDIUM]', 'low': '[LOW RISK]'}.get(sev, '[ALERT]')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            br = p.add_run(f'{badge}  ')
            br.bold = True
            br.font.size = Pt(9)
            br.font.color.rgb = col
            tr = p.add_run(alert.get('title', '') + ' — ')
            tr.bold = True
            tr.font.size = Pt(10)
            bdr = p.add_run(alert.get('body', ''))
            bdr.font.size = Pt(10)
    else:
        _para(doc, 'No significant risk alerts identified.', color=MID_GRAY)
    doc.add_paragraph()

    # SECTION 4 — CONDITIONS
    _heading(doc, '4. CONDITIONS — Loan Terms & Covenants')
    _kv_table(doc, [
        ('Decision',             decision.capitalize()),
        ('Recommended Amount',   f'Rs.{rec_loan} Cr' if rec_loan is not None else '—'),
        ('Interest Rate',        f'{interest}% per annum' if interest is not None else '—'),
        ('Tenor',                f'{tenor} months' if tenor is not None else '—'),
        ('Composite Risk Score', f'{risk_score} / 100'),
    ])
    if conditions:
        _heading(doc, 'Sanction Conditions', level=2)
        for cond in conditions:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.add_run(cond).font.size = Pt(10)
    doc.add_paragraph()

    # SECTION 5 — AI REASONING
    _heading(doc, '5. AI ANALYST REASONING')
    if reasoning:
        for para_text in reasoning.split('\n'):
            if para_text.strip():
                _para(doc, para_text.strip(), size=10, space_after=6)
    else:
        _para(doc, 'No reasoning available.', color=MID_GRAY)
    doc.add_paragraph()

    # FOOTER
    hr_p2 = doc.add_paragraph()
    pPr2 = hr_p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top')
    top2.set(qn('w:val'), 'single')
    top2.set(qn('w:sz'), '4')
    top2.set(qn('w:space'), '1')
    top2.set(qn('w:color'), 'CCCCCC')
    pBdr2.append(top2)
    pPr2.append(pBdr2)

    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_p.add_run(
        'This CAM is generated by the Intelli-Credit AI platform for internal use only. '
        'Final credit decisions are subject to Credit Committee review.'
    )
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = MID_GRAY
    disc_run.italic = True

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_run = gen_p.add_run(f'Intelli-Credit v2.1  •  Generated: {report_date} {report_time}')
    gen_run.font.size = Pt(8)
    gen_run.font.color.rgb = MID_GRAY

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()