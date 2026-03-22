"""
CAM (Credit Appraisal Memorandum) document generation endpoint.
Produces a professional Word (.docx) document structured around the Five Cs.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

router = APIRouter(prefix="/cam", tags=["CAM"])


# ── Request schema ──────────────────────────────────────────────────────────

class ScoreItem(BaseModel):
    score: Optional[float] = None
    weight_pct: Optional[float] = None

class RiskAlert(BaseModel):
    severity: Optional[str] = None
    title: Optional[str] = None
    body: Optional[str] = None

class YearlyRow(BaseModel):
    year: Optional[str] = None
    revenue: Optional[str] = None
    profit: Optional[str] = None
    debt: Optional[str] = None

class FinancialOverview(BaseModel):
    annual_revenue: Optional[str] = None
    annual_revenue_delta: Optional[str] = None
    annual_revenue_trend: Optional[str] = None
    net_profit: Optional[str] = None
    net_profit_margin: Optional[str] = None
    total_debt: Optional[str] = None
    de_ratio: Optional[str] = None
    gst_turnover: Optional[str] = None

class ScoreBreakdown(BaseModel):
    financial_health: Optional[ScoreItem] = None
    repayment_history: Optional[ScoreItem] = None
    collateral_coverage: Optional[ScoreItem] = None
    management_quality: Optional[ScoreItem] = None
    market_position: Optional[ScoreItem] = None

class SWOTItem(BaseModel):
    title: str = ""
    detail: str = ""

class SWOTData(BaseModel):
    strengths: List[SWOTItem] = []
    weaknesses: List[SWOTItem] = []
    opportunities: List[SWOTItem] = []
    threats: List[SWOTItem] = []

class CAMRequest(BaseModel):
    company_name: str = "—"
    sector: str = "—"
    location: str = "—"
    requested_loan_cr: Optional[float] = None
    recommended_loan_cr: Optional[float] = None
    decision: str = "—"
    risk_score: Optional[float] = None
    interest_rate_pct: Optional[float] = None
    tenor_months: Optional[int] = None
    conditions: List[str] = []
    risk_alerts: List[RiskAlert] = []
    reasoning: str = ""
    financial_overview: FinancialOverview = FinancialOverview()
    score_breakdown: ScoreBreakdown = ScoreBreakdown()
    yearly_trend: List[YearlyRow] = []
    primary_notes: List[Dict[str, Any]] = []
    swot: Optional[SWOTData] = None


# ── Helpers ──────────────────────────────────────────────────────────────────

DARK   = RGBColor(0x0F, 0x17, 0x2A)
SLATE  = RGBColor(0x64, 0x74, 0x8B)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
RED    = RGBColor(0xDC, 0x26, 0x26)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def _v(val, fallback="—"):
    """Return val if truthy, else fallback."""
    return str(val) if val is not None and str(val).strip() else fallback


def _fmt_cr(val):
    if val is None:
        return "—"
    return f"₹{val} Cr"


def _risk_label(score):
    if score is None:
        return "—"
    if score >= 70:
        return "HIGH RISK"
    if score >= 40:
        return "MEDIUM RISK"
    return "LOW RISK"


def _set_cell_shading(cell, hex_color: str):
    """Set cell background colour via XML."""
    shading = cell._element.get_or_add_tcPr()
    shading_el = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shading_el)


def _add_heading_band(doc: Document, text: str, color_hex: str = "2563EB"):
    """Add a coloured heading band as a single-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, color_hex)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    doc.add_paragraph()  # spacer


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    """Add a key-value table: Label | Value."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        # Label cell
        lc = tbl.rows[i].cells[0]
        lc.width = Cm(6)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = "Calibri"
        lr.font.size = Pt(9)
        lr.font.color.rgb = SLATE
        lr.bold = True

        # Value cell
        vc = tbl.rows[i].cells[1]
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = "Calibri"  # explicit — ensures ₹ renders correctly
        vr.font.size = Pt(9)
        vr.font.color.rgb = DARK

        # Alternate row shading
        if i % 2 == 0:
            _set_cell_shading(lc, "F8FAFC")
            _set_cell_shading(vc, "F8FAFC")

    doc.add_paragraph()  # spacer


def _score_str(item: Optional[ScoreItem]) -> str:
    if item is None or item.score is None:
        return "—"
    return f"{item.score}/100"


# ── Document builder ────────────────────────────────────────────────────────

def build_cam_docx(req: CAMRequest) -> io.BytesIO:
    doc = Document()

    # Set document default to Calibri — required for ₹ (U+20B9) to render correctly.
    # Without an explicit font name, python-docx uses theme fonts; depending on the
    # Office theme, ₹ may be substituted with '¹' (ordinal indicator).
    for style_name in ("Normal", "Default Paragraph Font"):
        try:
            doc.styles[style_name].font.name = "Calibri"
        except KeyError:
            pass

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    fo = req.financial_overview
    sb = req.score_breakdown
    now = datetime.now()
    report_date = now.strftime("%d %b %Y")
    report_time = now.strftime("%H:%M")

    rec_loan = _fmt_cr(req.recommended_loan_cr)
    req_loan = _fmt_cr(req.requested_loan_cr)
    decision = req.decision.capitalize() if req.decision else "—"

    interest = f"{req.interest_rate_pct}% per annum" if req.interest_rate_pct else "—"
    tenor = f"{req.tenor_months} months" if req.tenor_months else "—"

    # ── Title ────────────────────────────────────────────────────────────────
    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc = title_tbl.rows[0].cells[0]
    _set_cell_shading(tc, "0F172A")
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    tr = p.add_run("CREDIT APPRAISAL MEMORANDUM")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = WHITE

    sub = tc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.space_after = Pt(6)
    sr = sub.add_run(f"Intelli-Credit AI Platform  •  {report_date} {report_time} IST  •  CONFIDENTIAL")
    sr.font.size = Pt(8)
    sr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 1. Company Overview
    # ════════════════════════════════════════════════════════════════════════
    _add_heading_band(doc, "1. COMPANY OVERVIEW", "2563EB")
    _add_kv_table(doc, [
        ("Company Name",     _v(req.company_name)),
        ("Sector / Industry", _v(req.sector)),
        ("Location",          _v(req.location)),
        ("Loan Requested",    req_loan),
        ("Loan Recommended",  rec_loan),
        ("Decision",          decision),
    ])

    # ════════════════════════════════════════════════════════════════════════
    # 2. Financial Analysis
    # ════════════════════════════════════════════════════════════════════════
    _add_heading_band(doc, "2. FINANCIAL ANALYSIS", "16A34A")
    _add_kv_table(doc, [
        ("Annual Revenue",      _v(fo.annual_revenue)),
        ("Revenue Trend",       f"{_v(fo.annual_revenue_delta)} ({_v(fo.annual_revenue_trend)})" if fo.annual_revenue_delta else "—"),
        ("Net Profit",          _v(fo.net_profit)),
        ("Net Profit Margin",   _v(fo.net_profit_margin)),
        ("Total Debt",          _v(fo.total_debt)),
        ("Debt / Equity Ratio", _v(fo.de_ratio)),
        ("GST Turnover",        _v(fo.gst_turnover)),
    ])

    # Yearly trend table
    if req.yearly_trend:
        p = doc.add_paragraph()
        r = p.add_run("YEARLY TREND")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE

        tbl = doc.add_table(rows=1 + len(req.yearly_trend), cols=4)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, hdr in enumerate(["Year", "Revenue (Cr)", "Profit (Cr)", "Debt (Cr)"]):
            cell = tbl.rows[0].cells[j]
            _set_cell_shading(cell, "F1F5F9")
            pr = cell.paragraphs[0]
            hr = pr.add_run(hdr)
            hr.bold = True
            hr.font.size = Pt(8)
            hr.font.color.rgb = SLATE
        for i, row in enumerate(req.yearly_trend, start=1):
            for j, val in enumerate([_v(row.year), _v(row.revenue), _v(row.profit), _v(row.debt)]):
                cell = tbl.rows[i].cells[j]
                vr = cell.paragraphs[0].add_run(val)
                vr.font.size = Pt(8)
                vr.font.color.rgb = DARK
        doc.add_paragraph()

    # Score breakdown
    if sb:
        p = doc.add_paragraph()
        r = p.add_run("CREDIT SCORE BREAKDOWN")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE

        score_rows = [
            ("Financial Health",    sb.financial_health),
            ("Repayment History",   sb.repayment_history),
            ("Collateral Coverage", sb.collateral_coverage),
            ("Management Quality",  sb.management_quality),
            ("Market Position",     sb.market_position),
        ]
        tbl = doc.add_table(rows=len(score_rows), cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, item) in enumerate(score_rows):
            tbl.rows[i].cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
            score_val = _score_str(item)
            tbl.rows[i].cells[1].paragraphs[0].add_run(score_val).font.size = Pt(9)
            weight = f"{item.weight_pct}%" if item and item.weight_pct else "—"
            tbl.rows[i].cells[2].paragraphs[0].add_run(f"Weight: {weight}").font.size = Pt(8)
            if i % 2 == 0:
                for j in range(3):
                    _set_cell_shading(tbl.rows[i].cells[j], "F8FAFC")
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 3. Five Cs Credit Analysis
    # ════════════════════════════════════════════════════════════════════════
    _add_heading_band(doc, "3. FIVE Cs CREDIT ANALYSIS", "7C3AED")

    five_cs = [
        {
            "num": "1", "name": "Character",
            "color": "DBEAFE",
            "desc": "Willingness to repay based on credit history and management integrity.",
            "items": [
                ("Repayment History Score", _score_str(sb.repayment_history)),
                ("Management Quality Score", _score_str(sb.management_quality)),
            ],
        },
        {
            "num": "2", "name": "Capacity",
            "color": "DCFCE7",
            "desc": "Ability to repay based on cash flows and financial performance.",
            "items": [
                ("Financial Health Score", _score_str(sb.financial_health)),
                ("Annual Revenue", _v(fo.annual_revenue)),
                ("Net Profit Margin", _v(fo.net_profit_margin)),
            ],
        },
        {
            "num": "3", "name": "Capital",
            "color": "FEF3C7",
            "desc": "Financial strength and net worth of the borrower.",
            "items": [
                ("Debt / Equity Ratio", _v(fo.de_ratio)),
                ("Total Debt", _v(fo.total_debt)),
                ("GST Turnover", _v(fo.gst_turnover)),
            ],
        },
        {
            "num": "4", "name": "Collateral",
            "color": "FFE4E6",
            "desc": "Assets pledged as security against the loan.",
            "items": [
                ("Collateral Coverage Score", _score_str(sb.collateral_coverage)),
                ("Recommended Loan", rec_loan),
            ],
        },
        {
            "num": "5", "name": "Conditions",
            "color": "F1F5F9",
            "desc": "Macro environment, loan purpose, and imposed conditions.",
            "items": [
                ("Market Position Score", _score_str(sb.market_position)),
                ("Sector", _v(req.sector)),
                ("Loan Conditions", "; ".join(req.conditions) if req.conditions else "None"),
            ],
        },
    ]

    for c in five_cs:
        # C heading
        c_tbl = doc.add_table(rows=1, cols=1)
        c_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cc = c_tbl.rows[0].cells[0]
        _set_cell_shading(cc, c["color"])
        cp = cc.paragraphs[0]
        cp.space_before = Pt(2)
        cp.space_after = Pt(2)
        cr = cp.add_run(f"{c['num']}. {c['name']}")
        cr.bold = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = DARK

        # Description
        dp = doc.add_paragraph()
        dp.space_before = Pt(1)
        dr = dp.add_run(c["desc"])
        dr.italic = True
        dr.font.size = Pt(8)
        dr.font.color.rgb = SLATE

        # Items
        _add_kv_table(doc, c["items"])

    # ════════════════════════════════════════════════════════════════════════
    # 4. SWOT Analysis (only if data is present)
    # ════════════════════════════════════════════════════════════════════════
    has_swot = bool(
        req.swot and (
            req.swot.strengths or req.swot.weaknesses
            or req.swot.opportunities or req.swot.threats
        )
    )
    if has_swot:
        _add_heading_band(doc, "4. SWOT ANALYSIS", "7C3AED")
        swot_quadrants = [
            ("S — Strengths",     req.swot.strengths,     "16A34A"),
            ("W — Weaknesses",    req.swot.weaknesses,    "DC2626"),
            ("O — Opportunities", req.swot.opportunities, "2563EB"),
            ("T — Threats",       req.swot.threats,       "D97706"),
        ]
        for q_label, q_items, q_color in swot_quadrants:
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            qr = p.add_run(q_label)
            qr.bold = True
            qr.font.name = "Calibri"
            qr.font.size = Pt(9)
            try:
                qr.font.color.rgb = RGBColor(
                    int(q_color[0:2], 16),
                    int(q_color[2:4], 16),
                    int(q_color[4:6], 16),
                )
            except Exception:
                pass
            for item in q_items:
                bp = doc.add_paragraph(style="List Bullet")
                bp.space_before = Pt(0)
                title_r = bp.add_run(f"{item.title}: ")
                title_r.bold = True
                title_r.font.name = "Calibri"
                title_r.font.size = Pt(8)
                title_r.font.color.rgb = DARK
                detail_r = bp.add_run(item.detail)
                detail_r.font.name = "Calibri"
                detail_r.font.size = Pt(8)
                detail_r.font.color.rgb = SLATE
        doc.add_paragraph()

    # Offset subsequent section numbers if SWOT is present
    _sec = lambda n: str(n + (1 if has_swot else 0))

    # ════════════════════════════════════════════════════════════════════════
    # 5. Risk Assessment  (or 5 if no SWOT)
    # ════════════════════════════════════════════════════════════════════════
    _add_heading_band(doc, f"{_sec(4)}. RISK ASSESSMENT", "D97706")

    if req.risk_alerts:
        for alert in req.risk_alerts:
            sev = (alert.severity or "").upper()
            color = RED if sev == "HIGH" else AMBER if sev == "MEDIUM" else GREEN
            p = doc.add_paragraph()
            sr = p.add_run(f"[{sev}] ")
            sr.bold = True
            sr.font.size = Pt(9)
            sr.font.color.rgb = color
            tr = p.add_run(_v(alert.title))
            tr.bold = True
            tr.font.size = Pt(9)
            tr.font.color.rgb = DARK
            if alert.body:
                bp = doc.add_paragraph()
                bp.space_before = Pt(0)
                br = bp.add_run(_v(alert.body))
                br.font.size = Pt(8)
                br.font.color.rgb = SLATE
    else:
        p = doc.add_paragraph()
        r = p.add_run("No risk alerts identified in this assessment.")
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # Primary Due Diligence Notes
    # ════════════════════════════════════════════════════════════════════════
    if req.primary_notes:
        _add_heading_band(doc, f"{_sec(5)}. PRIMARY DUE DILIGENCE NOTES", "059669")
        type_labels = {
            "site_visit": "Site Visit",
            "management_interview": "Management Interview",
            "market_feedback": "Market Feedback",
            "operational": "Operational",
            "other": "Other",
        }
        for note in req.primary_notes:
            ntype = type_labels.get(note.get("type", ""), note.get("type", "Note"))
            p = doc.add_paragraph()
            lr = p.add_run(f"[{ntype}] ")
            lr.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = GREEN
            tr = p.add_run(str(note.get("text", "")))
            tr.font.size = Pt(9)
            tr.font.color.rgb = DARK
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # Final Recommendation
    # ════════════════════════════════════════════════════════════════════════
    section_num = _sec(6) if req.primary_notes else _sec(5)
    _add_heading_band(doc, f"{section_num}. FINAL RECOMMENDATION", "2563EB")
    _add_kv_table(doc, [
        ("Decision",              decision),
        ("Recommended Amount",    rec_loan),
        ("Interest Rate",         interest),
        ("Tenor",                 tenor),
        ("Composite Risk Score",  f"{_v(req.risk_score)}/100 — {_risk_label(req.risk_score)}"),
    ])

    if req.conditions:
        p = doc.add_paragraph()
        r = p.add_run("CONDITIONS PRECEDENT")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE
        for cond in req.conditions:
            bp = doc.add_paragraph(style="List Bullet")
            br = bp.add_run(cond)
            br.font.size = Pt(9)
            br.font.color.rgb = DARK

    if req.reasoning:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("AI REASONING")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE
        rp = doc.add_paragraph()
        rr = rp.add_run(req.reasoning)
        rr.font.size = Pt(8)
        rr.font.color.rgb = SLATE

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_tbl = doc.add_table(rows=1, cols=1)
    footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fc = footer_tbl.rows[0].cells[0]
    _set_cell_shading(fc, "F1F5F9")
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.space_before = Pt(4)
    fp.space_after = Pt(4)
    fr = fp.add_run(
        "This CAM is generated by Intelli-Credit AI and is for internal use only. "
        "Final credit decisions are subject to review by the Credit Committee as per applicable banking guidelines."
    )
    fr.font.size = Pt(7)
    fr.font.color.rgb = SLATE

    # Serialize to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Endpoint ────────────────────────────────────────────────────────────────

@router.post("/generate")
async def generate_cam(req: CAMRequest):
    """Generate a Word (.docx) Credit Appraisal Memorandum structured around the Five Cs."""
    buf = build_cam_docx(req)
    safe_name = req.company_name.replace(" ", "_").replace("/", "_")
    filename = f"CAM_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
